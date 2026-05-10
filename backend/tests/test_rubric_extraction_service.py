"""Rubric extraction — text → LLM → validated criteria."""

import io
import json

import pytest

from api.services.rubric_extraction_service import (
    RubricExtractionError,
    RubricExtractionService,
)
from db_models import Question


class _ScriptedLLM:
    """Test LLM that returns a fixed string from generate(). The conftest
    stub_llm fixture monkeypatches `get_llm_client`; this test installs its
    own scripted client at the module level."""

    def __init__(self, response: str):
        self._response = response
        self.calls = 0

    def score(self, system: str, user: str) -> dict:
        return {"score": 1, "reasoning": "n/a"}

    def generate(self, system: str, user: str) -> str:
        self.calls += 1
        return self._response


@pytest.fixture
def patch_llm(monkeypatch: pytest.MonkeyPatch):
    """Helper to install a scripted LLM for a single test."""
    def install(response: str) -> _ScriptedLLM:
        client = _ScriptedLLM(response)
        import api.services.rubric_extraction_service as mod
        monkeypatch.setattr(mod, "get_llm_client", lambda *a, **kw: client)
        return client
    return install


# ------------------------------------------------------------------ #
# Happy paths                                                        #
# ------------------------------------------------------------------ #


def test_extract_clean_json(patch_llm):
    response = json.dumps({
        "criteria": [
            {
                "name": "team_capability",
                "description": "Team experience and skills",
                "scale_min": 1,
                "scale_max": 3,
                "anchor_descriptions": {"1": "weak", "2": "ok", "3": "strong"},
                "feeding_question_keys": ["team"],
                "weighted_question_keys": [],
            },
        ],
    })
    patch_llm(response)

    criteria = RubricExtractionService().extract_from_text("Team capability is judged on Q: team")
    assert len(criteria) == 1
    c = criteria[0]
    assert c["name"] == "team_capability"
    assert c["scale_min"] == 1 and c["scale_max"] == 3
    assert c["anchor_descriptions"]["1"] == "weak"
    assert c["feeding_question_keys"] == ["team"]


def test_extract_strips_code_fence(patch_llm):
    """LLMs often wrap JSON in ```json fences."""
    response = "```json\n" + json.dumps({
        "criteria": [
            {
                "name": "clarity",
                "description": "Is the answer clear?",
                "scale_min": 1, "scale_max": 5,
                "anchor_descriptions": {
                    "1": "unclear", "2": "vague", "3": "ok", "4": "clear", "5": "crystal",
                },
            },
        ],
    }) + "\n```"
    patch_llm(response)

    criteria = RubricExtractionService().extract_from_text("Clarity 1-5...")
    assert criteria[0]["name"] == "clarity"
    assert len(criteria[0]["anchor_descriptions"]) == 5


def test_extract_strips_preamble(patch_llm):
    """LLMs sometimes add 'Here is the JSON:' before the JSON."""
    response = (
        "Here is the structured rubric you asked for:\n"
        + json.dumps({
            "criteria": [
                {
                    "name": "fit",
                    "description": "How well it fits",
                    "scale_min": 1, "scale_max": 3,
                    "anchor_descriptions": {"1": "low", "2": "med", "3": "high"},
                },
            ],
        })
        + "\n\nLet me know if you need anything else."
    )
    patch_llm(response)

    criteria = RubricExtractionService().extract_from_text("Fit on 1-3 scale")
    assert criteria[0]["name"] == "fit"


def test_extract_normalizes_missing_lists(patch_llm):
    """Missing feeding/weighted arrays default to []."""
    response = json.dumps({
        "criteria": [
            {
                "name": "x",
                "description": "X",
                "scale_min": 1, "scale_max": 3,
                "anchor_descriptions": {"1": "a", "2": "b", "3": "c"},
            },
        ],
    })
    patch_llm(response)

    criteria = RubricExtractionService().extract_from_text("anything")
    assert criteria[0]["feeding_question_keys"] == []
    assert criteria[0]["weighted_question_keys"] == []


# ------------------------------------------------------------------ #
# Validation errors                                                  #
# ------------------------------------------------------------------ #


def test_extract_rejects_empty_text():
    with pytest.raises(RubricExtractionError, match="Empty"):
        RubricExtractionService().extract_from_text("   ")


def test_extract_rejects_invalid_json(patch_llm):
    patch_llm("This is not JSON at all")
    with pytest.raises(RubricExtractionError, match="malformed JSON"):
        RubricExtractionService().extract_from_text("rubric text")


def test_extract_rejects_missing_criteria_key(patch_llm):
    patch_llm(json.dumps({"foo": "bar"}))
    with pytest.raises(RubricExtractionError, match="missing top-level 'criteria' key"):
        RubricExtractionService().extract_from_text("rubric text")


def test_extract_rejects_empty_criteria_list(patch_llm):
    patch_llm(json.dumps({"criteria": []}))
    with pytest.raises(RubricExtractionError, match="No criteria extracted"):
        RubricExtractionService().extract_from_text("rubric text")


def test_extract_rejects_criterion_missing_name(patch_llm):
    patch_llm(json.dumps({
        "criteria": [
            {"description": "x", "scale_min": 1, "scale_max": 3,
             "anchor_descriptions": {"1": "a", "2": "b", "3": "c"}},
        ],
    }))
    with pytest.raises(RubricExtractionError, match="missing a name"):
        RubricExtractionService().extract_from_text("text")


def test_extract_rejects_missing_anchors(patch_llm):
    patch_llm(json.dumps({
        "criteria": [
            {"name": "team", "description": "team",
             "scale_min": 1, "scale_max": 5,
             "anchor_descriptions": {"1": "weak", "5": "top"}},  # 2,3,4 missing
        ],
    }))
    with pytest.raises(RubricExtractionError, match="missing anchors"):
        RubricExtractionService().extract_from_text("text")


def test_extract_rejects_inverted_scale(patch_llm):
    patch_llm(json.dumps({
        "criteria": [
            {"name": "x", "description": "x",
             "scale_min": 5, "scale_max": 1,
             "anchor_descriptions": {"1": "a", "2": "b", "3": "c", "4": "d", "5": "e"}},
        ],
    }))
    with pytest.raises(RubricExtractionError, match="scale_max"):
        RubricExtractionService().extract_from_text("text")


# ------------------------------------------------------------------ #
# File extraction                                                    #
# ------------------------------------------------------------------ #


def test_extract_from_txt_file(patch_llm):
    patch_llm(json.dumps({
        "criteria": [
            {"name": "c1", "description": "c1",
             "scale_min": 1, "scale_max": 3,
             "anchor_descriptions": {"1": "a", "2": "b", "3": "c"}},
        ],
    }))

    text_bytes = "Some rubric text".encode("utf-8")
    criteria = RubricExtractionService().extract_from_file(text_bytes, "rubric.txt")
    assert criteria[0]["name"] == "c1"


def test_extract_from_md_file(patch_llm):
    patch_llm(json.dumps({
        "criteria": [
            {"name": "c1", "description": "c1",
             "scale_min": 1, "scale_max": 3,
             "anchor_descriptions": {"1": "a", "2": "b", "3": "c"}},
        ],
    }))

    md = "# Rubric\n\n## Criterion 1\nDescription".encode("utf-8")
    criteria = RubricExtractionService().extract_from_file(md, "rubric.md")
    assert len(criteria) == 1


def test_extract_from_docx_file(patch_llm):
    patch_llm(json.dumps({
        "criteria": [
            {"name": "c1", "description": "c1",
             "scale_min": 1, "scale_max": 3,
             "anchor_descriptions": {"1": "a", "2": "b", "3": "c"}},
        ],
    }))

    # Build a real docx in-memory using python-docx
    from docx import Document
    doc = Document()
    doc.add_paragraph("Criterion 1: clarity")
    doc.add_paragraph("Score 1-3 scale")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    criteria = RubricExtractionService().extract_from_file(buf.getvalue(), "rubric.docx")
    assert criteria[0]["name"] == "c1"


def test_extract_rejects_pdf():
    """PDF deliberately unsupported in v1; clear error message."""
    with pytest.raises(RubricExtractionError, match="PDF rubrics are not supported"):
        RubricExtractionService().extract_from_file(b"%PDF-1.4...", "rubric.pdf")


def test_extract_rejects_unknown_extension():
    with pytest.raises(RubricExtractionError, match="Unsupported file type"):
        RubricExtractionService().extract_from_file(b"x", "rubric.zip")


# ------------------------------------------------------------------ #
# Ref binding                                                        #
# ------------------------------------------------------------------ #


def _qs(*specs: tuple[str, str, int]) -> list[Question]:
    """Helper: list of Questions from (key, text, column_index) tuples."""
    return [
        Question(id=i + 1, project_id=1, key=k, text=t, column_index=col)
        for i, (k, t, col) in enumerate(specs)
    ]


def test_binding_resolves_pure_numeric_refs_via_column_index(patch_llm):
    """Rubric: 'columns 8 and 9'. Dataset key q1 came from spreadsheet column
    index 7 (col 8 in 1-based). The binder must map '8' → 'q1'."""
    patch_llm(json.dumps({
        "criteria": [
            {"name": "team", "description": "Team",
             "scale_min": 1, "scale_max": 3,
             "anchor_descriptions": {"1": "weak", "2": "ok", "3": "strong"},
             "feeding_question_keys": ["8", "9"],
             "weighted_question_keys": ["9"]},
        ],
    }))
    questions = _qs(
        ("q1", "Tim ima znanja", 7),  # col 8 (1-based)
        ("q2", "Iskustvo", 8),         # col 9
        ("q3", "Other", 9),
    )
    criteria = RubricExtractionService().extract_from_text("rubric", questions=questions)
    c = criteria[0]
    assert c["feeding_question_keys"] == ["q1", "q2"]
    assert c["weighted_question_keys"] == ["q2"]
    assert c["unbound_feeding_refs"] == []


def test_binding_resolves_q_prefixed_refs(patch_llm):
    patch_llm(json.dumps({
        "criteria": [
            {"name": "x", "description": "x",
             "scale_min": 1, "scale_max": 3,
             "anchor_descriptions": {"1": "a", "2": "b", "3": "c"},
             "feeding_question_keys": ["Q1", "q03"],
             "weighted_question_keys": []},
        ],
    }))
    questions = _qs(("q1", "T1", 0), ("q3", "T3", 2))
    criteria = RubricExtractionService().extract_from_text("rubric", questions=questions)
    assert criteria[0]["feeding_question_keys"] == ["q1", "q3"]


def test_binding_falls_back_to_text_substring(patch_llm):
    patch_llm(json.dumps({
        "criteria": [
            {"name": "x", "description": "x",
             "scale_min": 1, "scale_max": 3,
             "anchor_descriptions": {"1": "a", "2": "b", "3": "c"},
             "feeding_question_keys": ["team experience"],
             "weighted_question_keys": []},
        ],
    }))
    questions = _qs(
        ("q1", "Describe your team experience and skills", 0),
        ("q2", "Market size", 1),
    )
    criteria = RubricExtractionService().extract_from_text("rubric", questions=questions)
    assert criteria[0]["feeding_question_keys"] == ["q1"]


def test_binding_surfaces_unbound_refs(patch_llm):
    patch_llm(json.dumps({
        "criteria": [
            {"name": "x", "description": "x",
             "scale_min": 1, "scale_max": 3,
             "anchor_descriptions": {"1": "a", "2": "b", "3": "c"},
             "feeding_question_keys": ["8", "999", "nonsense"],
             "weighted_question_keys": ["999"]},
        ],
    }))
    questions = _qs(("q1", "First", 7))   # col 8
    criteria = RubricExtractionService().extract_from_text("rubric", questions=questions)
    c = criteria[0]
    assert c["feeding_question_keys"] == ["q1"]
    assert sorted(c["unbound_feeding_refs"]) == ["999", "nonsense"]
    # Weighted ref didn't bind, and even if it had it would have been dropped
    # because its key isn't in feeding.
    assert c["weighted_question_keys"] == []
    assert c["unbound_weighted_refs"] == ["999"]


def test_binding_skipped_when_no_questions_imported_yet(patch_llm):
    """Rubric upload before dataset import — leave refs untouched, don't flag."""
    patch_llm(json.dumps({
        "criteria": [
            {"name": "x", "description": "x",
             "scale_min": 1, "scale_max": 3,
             "anchor_descriptions": {"1": "a", "2": "b", "3": "c"},
             "feeding_question_keys": ["8", "9"],
             "weighted_question_keys": []},
        ],
    }))
    criteria = RubricExtractionService().extract_from_text("rubric")  # no questions
    assert criteria[0]["feeding_question_keys"] == ["8", "9"]
    assert criteria[0]["unbound_feeding_refs"] == []


def test_extraction_contract_for_question_plus_guidance_format(patch_llm):
    """Documents the contract for the prospect's rubric format:
    'Question :' line is the criterion description; 'Guidance for evaluators :'
    paragraphs are the per-score anchors (descending: first paragraph =
    scale_max). This test stubs the LLM with a CORRECT extraction; if the
    extractor's prompt regresses and starts confusing description with the
    score-3 anchor (or substituting global preamble for anchor text), the
    integration test against a real LLM should fail. Here we just lock in
    the downstream shape."""
    patch_llm(json.dumps({
        "criteria": [
            {
                "name": "team",
                "description": (
                    "U kojoj mjeri je tim sposoban iznijeti poduzetnički pothvat?"
                ),
                "scale_min": 1,
                "scale_max": 3,
                "anchor_descriptions": {
                    "3": "Tim ima znanja i iskustva za potrebe poduzetničkog pothvata te je izvjesno da može kvalitetno provesti poduzetnički pothvat.",
                    "2": "Tim nema sva potrebna znanja i iskustva za potrebe poduzetničkog pothvata.",
                    "1": "S obzirom na sastav tima, postoji rizik o uspješnoj provedbi poduzetničkog pothvata.",
                },
                "feeding_question_keys": ["8", "9", "10"],
                "weighted_question_keys": ["9"],
            },
        ],
    }))

    rubric_text = """
Scoring logic
3 points = meets criteria
2 points = has potential and partially meets criteria
1 point = it is a zero basically.

**Team**
Question : U kojoj mjeri je tim sposoban iznijeti poduzetnički pothvat?
Guidance for evaluators :
Tim ima znanja i iskustva za potrebe poduzetničkog pothvata te je izvjesno da može kvalitetno provesti poduzetnički pothvat.

Tim nema sva potrebna znanja i iskustva za potrebe poduzetničkog pothvata.

S obzirom na sastav tima, postoji rizik o uspješnoj provedbi poduzetničkog pothvata.

Evaluators use the answers in the following columns to arrive at score: 8, 9, 10. 9 is the most important question here.
""".strip()

    criteria = RubricExtractionService().extract_from_text(rubric_text)
    c = criteria[0]
    # Description is the criterion's QUESTION, not the score-3 anchor text.
    assert c["description"].startswith("U kojoj mjeri je tim sposoban")
    # Anchor for score 3 is the per-criterion guidance, NOT the global preamble.
    assert "Tim ima znanja i iskustva" in c["anchor_descriptions"]["3"]
    assert "meets criteria" not in c["anchor_descriptions"]["3"]
    # Lowest anchor reflects the per-criterion risk paragraph, not the legend.
    assert "rizik" in c["anchor_descriptions"]["1"]
    assert "Answer is not provided" not in c["anchor_descriptions"]["1"]


def test_binding_falls_back_to_q_key_when_column_index_unknown(patch_llm):
    """Synthetic data with no column_index (e.g. demo seeding) — numeric refs
    fall back to qN matching."""
    patch_llm(json.dumps({
        "criteria": [
            {"name": "x", "description": "x",
             "scale_min": 1, "scale_max": 3,
             "anchor_descriptions": {"1": "a", "2": "b", "3": "c"},
             "feeding_question_keys": ["3"],
             "weighted_question_keys": []},
        ],
    }))
    questions = [
        Question(id=1, project_id=1, key="q1", text="A", column_index=None),
        Question(id=2, project_id=1, key="q3", text="C", column_index=None),
    ]
    criteria = RubricExtractionService().extract_from_text("rubric", questions=questions)
    assert criteria[0]["feeding_question_keys"] == ["q3"]
